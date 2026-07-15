"""Document loading and text chunking utilities."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass(frozen=True)
class DocumentSection:
    source: str
    location: str
    text: str


@dataclass(frozen=True)
class TextChunk:
    source: str
    location: str
    chunk_index: int
    content: str


def iter_document_files(directory: Path) -> list[Path]:
    if not directory.exists():
        raise FileNotFoundError(f"Document directory does not exist: {directory}")
    return sorted(
        path
        for path in directory.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _read_pdf(path: Path, source: str) -> list[DocumentSection]:
    reader = PdfReader(path)
    sections: list[DocumentSection] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = _normalize_text(page.extract_text() or "")
        if text:
            sections.append(DocumentSection(source, f"page {page_number}", text))
    return sections


def _read_docx(path: Path, source: str) -> list[DocumentSection]:
    document = Document(path)
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            blocks.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    text = _normalize_text("\n".join(blocks))
    return [DocumentSection(source, "document", text)] if text else []


def _read_text(path: Path, source: str) -> list[DocumentSection]:
    text = _normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
    return [DocumentSection(source, "document", text)] if text else []


def read_document(path: Path, source_name: str | None = None) -> list[DocumentSection]:
    source = source_name or path.name
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported document type: {extension}. "
            f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    if extension == ".pdf":
        return _read_pdf(path, source)
    if extension == ".docx":
        return _read_docx(path, source)
    return _read_text(path, source)


def split_text(text: str, max_chars: int, overlap_chars: int) -> list[str]:
    text = _normalize_text(text)
    if not text:
        return []
    if len(text) <= max_chars:
        return [text]

    chunks: list[str] = []
    start = 0

    while start < len(text):
        target_end = min(start + max_chars, len(text))
        end = target_end

        if target_end < len(text):
            minimum_break = start + max_chars // 2
            for delimiter in ("\n\n", "\n", ". ", " "):
                position = text.rfind(delimiter, minimum_break, target_end)
                if position != -1:
                    end = position + len(delimiter)
                    break

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break

        next_start = max(end - overlap_chars, start + 1)
        start = next_start

    return chunks


def chunk_sections(
    sections: Iterable[DocumentSection],
    max_chars: int,
    overlap_chars: int,
) -> list[TextChunk]:
    output: list[TextChunk] = []
    chunk_index = 0

    for section in sections:
        for content in split_text(section.text, max_chars, overlap_chars):
            output.append(
                TextChunk(
                    source=section.source,
                    location=section.location,
                    chunk_index=chunk_index,
                    content=content,
                )
            )
            chunk_index += 1

    return output
